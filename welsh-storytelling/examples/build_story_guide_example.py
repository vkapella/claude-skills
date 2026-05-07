"""
Build the companion Story Guide for the BNY RAP deck.

This is the "scaffolding the audience never sees" — delivery coaching for the
presenter, built on Welsh's frameworks: Diamond facets, Repetition with
Variation, Curveball Question handling, the Power of "I Don't Know,"
Poise/Presence/Posture cues, and a Cut List.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# WWT Brand Colors
WWT_DARK_BLUE = RGBColor(0x1C, 0x00, 0x87)
WWT_LIGHT_BLUE = RGBColor(0x00, 0x86, 0xEA)
WWT_RED = RGBColor(0xEE, 0x28, 0x2A)
NAVY = RGBColor(0x1D, 0x1E, 0x48)
ORANGE = RGBColor(0xFB, 0x55, 0x0E)
BODY = RGBColor(0x26, 0x26, 0x26)
GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY_BG = "F5F7FA"
LIGHT_BLUE_TINT = "E6F3FD"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Set base font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = BODY

# ============================================================================
# Helpers
# ============================================================================
def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_horizontal_rule(doc, color="0086EA"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def cover_title(text, size=28, color=WWT_DARK_BLUE, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = 'Calibri'
    return p

def eyebrow(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(9)
    run.font.color.rgb = ORANGE
    run.font.bold = True
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    return p

def section_heading(text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = WWT_DARK_BLUE if level == 1 else NAVY
    run.font.bold = True
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def body_para(text, italic=False, size=11, color=None, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold
    run.font.name = 'Calibri'
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p

def labeled_para(label, text, label_color=WWT_LIGHT_BLUE):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}  ")
    r1.font.bold = True
    r1.font.color.rgb = label_color
    r1.font.size = Pt(11)
    r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = BODY
    r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.size = Pt(11)
    run.font.color.rgb = BODY
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(3)
    return p

def callout_box(label, text, bg=LIGHT_BLUE_TINT, accent=WWT_LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_bg(cell, bg)

    # Add a left border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), '0086EA')
    tcBorders.append(left)
    tcPr.append(tcBorders)

    # Padding
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top','120'),('left','180'),('bottom','120'),('right','180')]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), val)
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

    # Content
    cell.text = ''
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(label.upper())
    r1.font.size = Pt(9)
    r1.font.color.rgb = accent
    r1.font.bold = True
    r1.font.name = 'Calibri'
    p1.paragraph_format.space_after = Pt(2)

    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(11)
    r2.font.italic = True
    r2.font.color.rgb = BODY
    r2.font.name = 'Calibri'

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# ============================================================================
# COVER
# ============================================================================
eyebrow("Story guide  ·  Companion to the slide deck")
cover_title("11 sites. 6 months. One decision.", size=26)
add_horizontal_rule(doc)
body_para("Delivery coaching for the BNY RAP Microbranch program presentation. Built on Michael Welsh's storytelling frameworks from The Backstory on Storytelling — applied to your deck.",
          italic=True, color=GRAY)

p = doc.add_paragraph()
p.add_run("WWT  ·  Prepared for the BNY COO meeting  ·  March 2026").font.size = Pt(10)
p.runs[0].font.color.rgb = GRAY

doc.add_paragraph()

# ============================================================================
# 1. CORE MESSAGE
# ============================================================================
section_heading("1. The core message", level=1)
body_para("If the COO leaves the room remembering one sentence, this is the sentence:")

callout_box(
    "Core message",
    "Cisco SD-WAN costs $1.4M more than Meraki — but it's the only option that fixes all three documented root causes AND aligns with the broader Branch Refresh. We're not buying eleven branches. We're buying the pattern for the next two hundred."
)

body_para("Everything else in the deck — the architecture, the conformance table, the cost numbers — is in service of this one sentence. When in doubt during Q&A, return to this.")

# ============================================================================
# 2. THE DIAMOND
# ============================================================================
section_heading("2. The Diamond", level=1)
body_para("Welsh's framework: every core message is a diamond with five facets. Same truth, different angles. Use whichever facet catches the light for the question being asked.",
          italic=True, color=GRAY)

# Diamond table
diamond_table = doc.add_table(rows=5, cols=2)
diamond_table.autofit = False
diamond_table.columns[0].width = Inches(1.6)
diamond_table.columns[1].width = Inches(4.9)

facets = [
    ("The Point",       "We recommend Cisco SD-WAN for the eleven RAP sites — full stop."),
    ("The Context",     "The previous build broke in three documented ways. The new design has to fix all three."),
    ("The Evidence",    "12 of 12 design requirements met fully. Meraki misses three — including end-to-end segmentation and forward alignment."),
    ("The Application", "BNY operates one network, not two. The 2027 rollout uses the same pattern. The NOC stops being reactive."),
    ("The Call to Action", "COO sign-off this week. Services contract by May 1. Eleven sites turnkey by end of Q3."),
]

for i, (label, text) in enumerate(facets):
    row = diamond_table.rows[i]
    row.cells[0].text = ''
    row.cells[1].text = ''

    set_cell_bg(row.cells[0], LIGHT_BLUE_TINT)

    p1 = row.cells[0].paragraphs[0]
    r1 = p1.add_run(label)
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r1.font.name = 'Calibri'

    p2 = row.cells[1].paragraphs[0]
    r2 = p2.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = BODY
    r2.font.name = 'Calibri'

    # Cell padding
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top','100'),('left','140'),('bottom','100'),('right','140')]:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), val)
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)

doc.add_paragraph()

# ============================================================================
# 3. REPETITION WITH VARIATION
# ============================================================================
section_heading("3. Repetition with Variation library", level=1)
body_para("Welsh's principle: tell the same truth three different ways. Pick 2-3 of these and weave them across the talk so the message lands without sounding like a broken record.",
          italic=True, color=GRAY)

section_heading("Three ways to land \"why Cisco, not Meraki\"", level=2)
labeled_para("Technical",  "\"One architecture, branch to data center.\"")
labeled_para("Financial",  "\"Pay $1.4M now, or pay it twice in eighteen months.\"")
labeled_para("Strategic",  "\"Meraki gets us through the door. Cisco SD-WAN gets us through the next decade.\"")

section_heading("Three ways to land the cost framing", level=2)
labeled_para("Honest",     "\"$1.4M more is a real number. Here's what it buys.\"")
labeled_para("Operational", "\"The cheaper option is two networks. The recommended option is one.\"")
labeled_para("Future-facing", "\"We're not buying eleven branches. We're buying the pattern for two hundred.\"")

section_heading("Three ways to close", level=2)
labeled_para("Rhythmic",      "\"One decision, eleven sites, one pattern.\"")
labeled_para("Action-forward", "\"Sign this week, ship by September.\"")
labeled_para("Strategic",      "\"Today's decision is what gets reused two hundred times.\"")

doc.add_paragraph()

# ============================================================================
# 4. ARC MAP
# ============================================================================
section_heading("4. Arc map  ·  How time gets spent", level=1)
body_para("Welsh's math: 60-minute meeting = ~47 minutes of material maximum. This deck is built for a 30-minute slot, leaving real room for discussion.",
          italic=True, color=GRAY)

arc_table = doc.add_table(rows=7, cols=4)
arc_table.autofit = False
arc_table.columns[0].width = Inches(0.5)
arc_table.columns[1].width = Inches(1.3)
arc_table.columns[2].width = Inches(2.7)
arc_table.columns[3].width = Inches(0.8)

# Header
hdr = arc_table.rows[0]
for i, label in enumerate(["#", "Beat", "What happens", "Time"]):
    cell = hdr.cells[i]
    set_cell_bg(cell, "1D1E48")
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'

beats = [
    ("1", "Title",      "Open. Set the promise: one decision, today.",        "0:30"),
    ("2", "Setup",      "Three documented failures. Establish the constraint.", "3:00"),
    ("3", "What",       "Core message. The Diamond's top facet.",              "2:00"),
    ("4", "How",        "Three requirements that decide it. Walk the table.",  "5:00"),
    ("5", "Why",        "Cost framed honestly. The strategic frame.",          "4:00"),
    ("6", "Magic Wand", "The ask. The picture of success. SHUT UP.",           "1:30"),
]

for i, (num, beat, what, time) in enumerate(beats, start=1):
    row = arc_table.rows[i]
    for j, val in enumerate([num, beat, what, time]):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY
        run.font.name = 'Calibri'
        if j == 1: run.font.bold = True

body_para("\nTotal slide time: ~16 minutes. Leaves ~14 minutes for Q&A inside a 30-minute slot. If asked to fill 60 minutes, do not add slides — extend Q&A.",
          italic=True, color=GRAY)

# ============================================================================
# 5. CURVEBALL PREP
# ============================================================================
section_heading("5. Curveball question prep", level=1)
body_para("Welsh's six techniques for handling hard questions: don't react immediately, acknowledge, intentional silence, characterize and categorize, disambiguate and redirect, complete disassembly. The questions below are the ones the COO is most likely to throw — with drafted answers using those techniques.",
          italic=True, color=GRAY)

curveballs = [
    {
        "q": "$1.4M is a lot of money. Why shouldn't we just go Meraki?",
        "tech": "Characterize and categorize",
        "a": "\"That's the right question to push on. Let me restate what I'm hearing: you're asking whether the strategic value of architectural alignment justifies a 90% cost premium on this specific spend. It does — but only if you believe we're going to roll out branches at volume in 2027. If you don't believe that, Meraki is the right answer. Do you believe it?\""
    },
    {
        "q": "Can we do Meraki at some sites and Cisco at others?",
        "tech": "Complete disassembly",
        "a": "\"Operating two networks at the same time is exactly what slide 2 said broke last time. The fragmented management plane was one of three documented root causes. A hybrid Cisco/Meraki approach gives us the cost of Cisco AND the operational pain of Meraki. I'd rather we pick one and commit.\""
    },
    {
        "q": "Why not wait six months and see what Branch Refresh actually picks?",
        "tech": "Disambiguate and redirect",
        "a": "\"Help me understand the concern — is it about technical risk, or about budget timing? If it's technical: Branch Refresh has already committed to Cisco SD-WAN. If it's budget: waiting costs us the eleven sites that need to be live by end of Q3, plus the 182-day SD-WAN lead time means waiting six months means deploying in 2027.\""
    },
    {
        "q": "What if the broader Branch Refresh changes direction?",
        "tech": "Honest \"I don't know\"",
        "a": "\"Honestly — I can't guarantee what Branch Refresh will look like in two years. What I can tell you is that today's pick aligns with today's plan. If Refresh shifts to a different vendor, we'd have a forklift in front of us either way. Cisco SD-WAN is the lower-risk bet given what's known now.\""
    },
    {
        "q": "Has WWT done this before at our scale?",
        "tech": "Acknowledge, then evidence",
        "a": "\"Yes — we've delivered automation and refresh programs for eight of the top ten US financials. The eleven sites here are the pilot for the WWT Branch-in-a-Box program at BNY, which is exactly designed to be the repeatable pattern for your 2027 volume. Happy to walk through specific reference architectures if useful.\""
    },
    {
        "q": "What's the rollback plan if the first sites struggle?",
        "tech": "Acknowledge, then concrete plan",
        "a": "\"Every site has a documented rollback procedure in the runbook — the existing connectivity stays in place until the new path is validated and signed off. The first three sites are intentionally low-risk geographies so we can find and fix issues before scaling. If the first three don't go cleanly, we pause and re-plan rather than push through.\""
    },
    {
        "q": "What's the one thing that could blow this timeline up?",
        "tech": "Honest \"I don't know\" + power move",
        "a": "\"The 182-day Cisco SD-WAN lead time is the longest pole in the tent. If the PO slips past mid-April, the September date slips. Beyond that — customs and import delays at the international sites are the next risk. We have a concierge service for that, but it's not zero risk. I won't promise it can't happen.\""
    },
]

for cb in curveballs:
    # Question
    p = doc.add_paragraph()
    r = p.add_run(f"Q.  {cb['q']}")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = WWT_DARK_BLUE
    r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True

    # Technique
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Technique:  {cb['tech']}")
    r2.font.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = ORANGE
    r2.font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.keep_with_next = True

    # Answer
    p3 = doc.add_paragraph()
    r3 = p3.add_run(cb['a'])
    r3.font.size = Pt(11)
    r3.font.color.rgb = BODY
    r3.font.name = 'Calibri'
    p3.paragraph_format.left_indent = Inches(0.3)
    p3.paragraph_format.space_after = Pt(6)

# ============================================================================
# 6. THE "I DON'T KNOW" LIST
# ============================================================================
section_heading("6. Topics where \"I don't know\" is the right answer", level=1)
body_para("Welsh's power move: admitting uncertainty turns vendors into partners. Have these queued up so you can deploy them without scrambling.",
          italic=True, color=GRAY)

idk_items = [
    ("Specific OEM hardware lifecycles past 2028", "We can commit to current Cisco roadmap; beyond 36 months is speculation."),
    ("Exact go-live date for any specific international site", "Customs and local readiness vary. We give a range, not a date, until the site is in flight."),
    ("Whether Branch Refresh selections will change", "Not our call. We design for the current state and stay flexible."),
    ("Total cost of the 2027 scaled rollout", "Depends on volume and OEM pricing then. We have a model, not a number."),
]

for topic, note in idk_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"·  {topic}  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r1.font.name = 'Calibri'
    r2 = p.add_run(note)
    r2.font.size = Pt(10)
    r2.font.italic = True
    r2.font.color.rgb = GRAY
    r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

# ============================================================================
# 7. OPENING & CLOSING LINES
# ============================================================================
section_heading("7. Opening and closing lines  ·  Memorize these verbatim", level=1)

eyebrow("First 30 seconds")
callout_box(
    "Opening (verbatim)",
    "\"Thanks for the time. This deck has one purpose — to help us land on the right network for the eleven new RAP sites we're standing up over the next six months. There's a recommendation at the end. Before that, I want to make sure we agree on what broke last time, because that's what's shaping the choice.\""
)

eyebrow("Last 30 seconds")
callout_box(
    "Closing (verbatim)",
    "\"One decision today. Eleven sites by September. The pattern for everything that follows. — What questions do you have?\"  Then SHUT UP. Three-second rule. Make them speak first.",
    bg="FFF4E6", accent=ORANGE
)

# ============================================================================
# 8. DELIVERY NOTES
# ============================================================================
section_heading("8. Delivery notes  ·  Poise, presence, posture", level=1)

section_heading("Before you walk in", level=2)
bullet("Box breathing for one minute: inhale 4, hold 4, exhale 4, hold 4. Repeat four times.")
bullet("Read the room before you start. Two leaning back? Slow down. Phones out? Skip to the cost slide.")
bullet("Have the appendix bookmarked. You will need slide 3 (12-requirement table) and slide 7 (cost detail) on demand.")

section_heading("During delivery", level=2)
bullet("Three-Second Rule: pause three full seconds before clicking from slide to slide. The slide change is your friend, not your enemy.")
bullet("Eye contact: hold each person's gaze for one complete thought (3-5 seconds), then move on. Not the screen, not the back wall.")
bullet("Move. Don't hide behind the laptop. Step toward the audience when you make the cost case on slide 5.")
bullet("If you stumble, do not apologize. Pause, breathe, restart the sentence. Apologies amplify the stumble.")

section_heading("Where to slow down", level=2)
bullet("Slide 2, after listing the three failures. Let \"three root causes\" sit in the air for two seconds.")
bullet("Slide 3, after reading the navy box aloud. This is the message. Don't paraphrase it.")
bullet("Slide 5, the bottom line. Read it slowly, almost too slowly: \"Meraki saves money on these eleven sites. Cisco SD-WAN saves money on the next two hundred.\"")
bullet("Slide 6, after the closing line. SILENCE. Do not fill it.")

# ============================================================================
# 9. THE CUT LIST
# ============================================================================
section_heading("9. The cut list  ·  What's in the appendix and why it's not in the front", level=1)
body_para("Welsh: \"Make people pitch why their content deserves inclusion.\" These items were considered for the front of the deck and moved to the appendix or cut entirely. If asked, you can pull them up — but they don't belong in the story.",
          italic=True, color=GRAY)

cut_table = doc.add_table(rows=7, cols=2)
cut_table.autofit = False
cut_table.columns[0].width = Inches(2.5)
cut_table.columns[1].width = Inches(4.0)

cuts = [
    ("Content", "Why it's not in the front"),
    ("12-requirement conformance table", "Surfaces 3 in slide 4. Full table in appendix for technical Q&A."),
    ("Detailed program timeline / Gantt", "COO doesn't need it. Program manager does. Appendix."),
    ("RACI matrix", "Operational detail. Belongs in the SOW conversation, not this one."),
    ("Branch site list with topology diagrams", "Geography is reassuring but not decision-relevant. Appendix."),
    ("WWT capability slides (ATC, Integration Services, Global Delivery, Storefront)", "These are credentials, not arguments. Useful if asked \"can WWT do this\" — not before."),
    ("Intelligent Lifecycle Management diagram", "Beautiful framework. Wrong audience. The COO wants the decision, not the methodology."),
]

for i, (a, b) in enumerate(cuts):
    row = cut_table.rows[i]
    row.cells[0].text = ''
    row.cells[1].text = ''

    if i == 0:
        set_cell_bg(row.cells[0], "1D1E48")
        set_cell_bg(row.cells[1], "1D1E48")

    for j, val in enumerate([a, b]):
        cell = row.cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if i == 0:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        else:
            run.font.color.rgb = BODY if j == 1 else NAVY
            run.font.bold = (j == 0)

# ============================================================================
# FOOTER NOTE
# ============================================================================
doc.add_paragraph()
add_horizontal_rule(doc)
p = doc.add_paragraph()
r = p.add_run("This Story Guide applies frameworks from The Backstory on Storytelling by Michael Welsh: the 5 Slide Rule, the Setup-What-How-Why-Magic Wand arc, Designing for Diamonds, Repetition with Variation, and the Curveball Question techniques. The deck is the storyboard. This guide is the script.")
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = GRAY
r.font.name = 'Calibri'

# ============================================================================
# Save
# ============================================================================
output_path = "/home/claude/output/BNY_RAP_Story_Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
